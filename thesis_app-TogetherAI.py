import streamlit as st
import time
import json
from datetime import datetime
import random
import logging
import os
from together import Together
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
import re

# Configure logging
log_dir = "logs"
if not os.path.exists(log_dir):
    os.makedirs(log_dir)

log_file = os.path.join(log_dir, f"thesis_generator_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log")
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(log_file),
        logging.StreamHandler()
    ]
)

# Initialize OpenRouter client
client = Together(api_key="58ff5e76928a9cdc29e9dfcce0d4b8086178282df9850070cff41b7cfad43231")

# Configure the page
st.set_page_config(
    page_title="Thesis Generator",
    page_icon="📚",
    layout="wide"
)

# Create a container for the sticky header
header_container = st.container()

with header_container:
    # Title and description
    st.title("📚 Academic Thesis Generator")
    st.markdown("Generate comprehensive thesis sections using AI models with professional academic writing standards.")
    st.markdown("---")  # Add a separator line

# Sidebar for configuration
st.sidebar.header("Configuration")

# Input fields with default values
topic = st.sidebar.text_input(
    "Research Topic", 
    value="Sentiment Analysis using Pretrained models",
    help="Enter your research topic here"
)

methodology = st.sidebar.text_area(
    "Methodology", 
    value="Models used are BERT, GPT, XLNet.",
    help="Describe your research methodology"
)

model_name = st.sidebar.selectbox(
    "Select Model",
    [
        "deepseek-ai/DeepSeek-R1-Distill-Llama-70B-free"
    ],
    index=0
)

# Thesis sections
sections = [
    'Background', 'Problem Statement', 'Research Questions', 'Aims and Objective',
    'Significance of the study', 'Inscope of the study', 'Out of Scope',
    'Reasons for Defining the Scope', 'Business Understanding', 'Data Selection',
    'Data Preprocessing', 'Models', 'Implementation workflow', 'Model Evaluation',
    'Hardware Resources Required', 'Software Resources Required'
]

# Section selection
st.sidebar.subheader("Section Selection")
selected_sections = st.sidebar.multiselect(
    "Choose sections to generate",
    sections,
    default=sections[:3],  # Default to first 3 sections
    help="Select which sections you want to generate"
)

# Main content area
col1, col2 = st.columns([2, 1])

with col1:
    st.header("Generated Thesis Content")
    
    # Display current configuration
    st.info(f"**Topic:** {topic}\n\n**Methodology:** {methodology}")
    
    # Generate button
    generate_btn = st.button("🚀 Generate Thesis Sections", type="primary")
    
    # Progress tracking
    if 'generation_progress' not in st.session_state:
        st.session_state.generation_progress = {}
    
    if 'conversation_history' not in st.session_state:
        st.session_state.conversation_history = [
            {
                "role": "system",
                "content": "You are an expert academic writer with a PhD-level understanding of Data Science. Respond with professional, well-researched, plagiarism-free content suitable for a thesis. Use clear, formal academic English, and avoid overly casual or generic phrasing."
            }
        ]

with col2:
    pass  # Empty column for layout balance

def get_section_content(model, messages, max_retries=3, base_delay=120):
    """
    Function to get content from OpenRouter API with retry logic
    """
    for attempt in range(max_retries):
        try:
            logging.info(f"Attempting API call with model: {model} (Attempt {attempt + 1}/{max_retries})")
            completion = client.chat.completions.create(
                extra_body={},
                model=model,
                messages=messages
            )
            logging.info(f"Successfully generated content with model: {model}")
            return completion
        except Exception as e:
            error_message = str(e)
            logging.error(f"API Error: {error_message}")
            if "429" in error_message and "Rate limit exceeded" in error_message:
                st.error("""
                ⚠️ **Free API Quota Exceeded**
                
                The free quota for the selected model has been utilized. To continue using the service, please:
                
                1. Contact our service team at:
                   - Email: info@writerighthesis.in
                   - Phone: +91 8350947394 / +91 9932137822
                
                2. Or visit our website to upgrade to a paid version:
                   [www.writerighthesis.in](https://www.writerighthesis.in/)
                """)
                return None
            elif "429" in error_message and attempt < max_retries - 1:
                # Calculate delay with exponential backoff and jitter
                delay = base_delay * (2 ** attempt) + random.uniform(0, 30)
                logging.warning(f"Rate limit hit. Waiting {int(delay)} seconds before retry...")
                st.warning(f"Rate limit hit. Waiting {int(delay)} seconds before retry...")
                time.sleep(delay)
                continue
            else:
                st.error(f"API Error: {error_message}")
                return None

def clean_content(content):
    """
    Remove content between <think> tags
    """
    return re.sub(r'<think>.*?</think>', '', content, flags=re.DOTALL)

def create_word_document(content_dict, topic, methodology):
    """
    Create a Word document with specified formatting
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    
    # Add title
    title = doc.add_heading(topic, level=1)
    title.bold = True
    
    # Add methodology
    doc.add_paragraph(f"Methodology: {methodology}")
    doc.add_paragraph()  # Add spacing
    
    # Add sections
    for section, content in content_dict.items():
        # Clean content by removing think tags
        cleaned_content = clean_content(content)
        
        # Add section heading
        heading = doc.add_heading(section, level=1)
        heading.bold = True
        
        # Add section content
        doc.add_paragraph(cleaned_content)
        doc.add_paragraph()  # Add spacing between sections
    
    return doc

# Generation logic
if generate_btn:
    if not selected_sections:
        st.warning("Please select at least one section to generate.")
        logging.warning("No sections selected for generation")
    else:
        st.success("Starting thesis generation...")
        logging.info(f"Starting thesis generation for topic: {topic}")
        logging.info(f"Selected sections: {selected_sections}")
        logging.info(f"Selected model: {model_name}")
        
        # Initialize progress tracking
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        # Initialize generated content storage
        if 'generated_content' not in st.session_state:
            st.session_state.generated_content = {}
            logging.info("Initialized new generated content storage")
        
        # Initialize conversation history if not exists
        if 'conversation_history' not in st.session_state:
            st.session_state.conversation_history = [
                {
                    "role": "system",
                    "content": "You are an expert academic writer with a PhD-level understanding of Data Science. Respond with professional, well-researched, plagiarism-free content suitable for a thesis. Use clear, formal academic English, and avoid overly casual or generic phrasing."
                }
            ]
            logging.info("Initialized new conversation history")
        
        # Generate content for each selected section
        for i, section in enumerate(selected_sections):
            status_text.text(f"Generating {section}...")
            logging.info(f"Starting generation for section: {section}")
            
            user_prompt = f"Write the {section} section for a thesis on '{topic}'. The methodology includes: {methodology}. Avoid plagiarism. Use a human-like academic tone."
            
            # Add user message to conversation history
            st.session_state.conversation_history.append({
                "role": "user",
                "content": user_prompt
            })
            
            try:
                # Call the model with updated conversation history and retry logic
                completion = get_section_content(model_name, st.session_state.conversation_history)
                
                # If completion is None, it means rate limit was exceeded
                if completion is None:
                    break  # Exit the loop if rate limit is exceeded
                
                # Extract and store assistant message
                if completion and completion.choices:
                    assistant_message = completion.choices[0].message.content
                    st.session_state.conversation_history.append({
                        "role": "assistant",
                        "content": assistant_message
                    })
                    
                    # Store generated content
                    st.session_state.generated_content[section] = assistant_message
                    logging.info(f"Successfully generated content for section: {section}")
                    
                    # Update progress
                    progress = (i + 1) / len(selected_sections)
                    progress_bar.progress(progress)
                    
                    # Display generated section
                    with st.expander(f"📄 {section}", expanded=True):
                        st.write(clean_content(assistant_message))
                    
                else:
                    error_msg = f"Error generating content for {section}"
                    st.error(error_msg)
                    logging.error(error_msg)
                
                # Add constant delay with randomness between requests
                if i < len(selected_sections) - 1:  # Don't delay after the last section
                    delay = 60 + random.uniform(0, 30)  # 60 seconds base + up to 30 seconds random
                    logging.info(f"Waiting {int(delay)} seconds before next request")
                    time.sleep(delay)
                    
            except Exception as e:
                error_msg = f"Error generating {section}: {str(e)}"
                st.error(error_msg)
                logging.error(error_msg)
        
        # Only show completion message if we didn't hit rate limit
        if completion is not None:
            status_text.text("Generation complete!")
            st.success("✅ All sections generated successfully!")
            logging.info("Thesis generation completed successfully")
            
            # Create and offer Word document download
            if st.session_state.generated_content:
                doc = create_word_document(st.session_state.generated_content, topic, methodology)
                
                # Save the document to a temporary file
                temp_docx = "WRT.docx"
                doc.save(temp_docx)
                
                # Offer the document for download
                with open(temp_docx, "rb") as file:
                    st.download_button(
                        label="📥 Download Word Document",
                        data=file,
                        file_name="WRT.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
                # Clean up the temporary file
                os.remove(temp_docx)

# Display previously generated content if available
if 'generated_content' in st.session_state and st.session_state.generated_content:
    st.header("📋 Generated Sections")
    
    for section, content in st.session_state.generated_content.items():
        with st.expander(f"📄 {section}"):
            st.write(content)
            
            # Individual section actions
            col1, col2 = st.columns(2)
            with col1:
                if st.button(f"📋 Copy {section}", key=f"copy_{section}"):
                    st.code(content)
            with col2:
                if st.button(f"🔄 Regenerate {section}", key=f"regen_{section}"):
                    st.info(f"Regenerating {section}... (Feature to be implemented)")

# Footer
st.markdown("---")
st.markdown("""
### Contact Us & Let's Get Started

**Powered by WriteRightThesis**  
[www.writerighthesis.in](https://www.writerighthesis.in/)

**Contact Numbers:**  
+91 8350947394 / +91 9932137822

**Email:**  
info@writerighthesis.in

© 2025 by WriteRightThesis  
*Proudly Made in India*
""")

# Add log file download button in the sidebar at the bottom
st.sidebar.markdown("---")
st.sidebar.subheader("Logs")
if os.path.exists(log_file):
    with open(log_file, 'r') as f:
        log_content = f.read()
        st.sidebar.download_button(
            "📥 Download Logs",
            log_content,
            file_name=f"thesis_generator_logs_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log",
            mime="text/plain"
        )
